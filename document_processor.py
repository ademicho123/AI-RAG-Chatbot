from langchain.document_loaders import TextLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document  # Import the Document class
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

class DocumentProcessor:
    def __init__(self, file_path):
        self.file_path = file_path

    def load_and_split(self):
        documents = []
        if self.file_path.endswith('.txt'):
            loader = TextLoader(self.file_path)
            documents = loader.load()
        elif self.file_path.endswith('.pdf'):
            reader = PdfReader(self.file_path)
            for page in reader.pages:
                # Create a Document object for each page
                documents.append(Document(page_content=page.extract_text()))
        elif self.file_path.endswith('.doc') or self.file_path.endswith('.docx'):
            doc = DocxDocument(self.file_path)
            for para in doc.paragraphs:
                # Create a Document object for each paragraph
                documents.append(Document(page_content=para.text))
        else:
            raise ValueError("Unsupported file type")

        splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        texts = splitter.split_documents(documents)
        return texts

if __name__ == "__main__":
    processor = DocumentProcessor("data/sample_text.txt")
    texts = processor.load_and_split()
    print(f"Processed {len(texts)} text chunks")